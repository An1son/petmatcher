"""
Insert a Testing & Debugging Methodology table before the Validation Results Table
in FinalReport_Petmatcher_INFO4290.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import os

doc_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'FinalReport_Petmatcher_INFO4290.docx')
doc = Document(doc_path)

# Testing data for each requirement
testing_data = [
    [
        "User Authentication & Role Management",
        "Manual functional testing",
        "Registered new accounts as both Adopter and Shelter roles; verified login/logout flows, session persistence across page reloads, and confirmed that role-based routing correctly restricts access (e.g., Shelter dashboard inaccessible to Adopter accounts)."
    ],
    [
        "Shelter Dashboard \u2013 Pet Management",
        "CRUD testing via Shelter account",
        "Logged in as a Shelter user and performed full create, read, update, and delete operations on pet listings. Verified that changes persisted in the database and reflected immediately in the browse view for adopters."
    ],
    [
        "Adopter Preference Profiling",
        "Form validation and persistence testing",
        "Submitted preference forms with valid, invalid, and edge-case inputs (e.g., empty fields, custom species text). Confirmed Zod schema validation caught malformed entries and that saved preferences matched what was stored in the profiles table."
    ],
    [
        "Swipe-based Discovery Interface",
        "Manual gesture and animation testing",
        "Tested swipe interactions across Chrome, Firefox, and Safari on both desktop (mouse drag) and mobile (touch). Verified that swipe velocity/distance thresholds triggered correct Like/Pass actions, and that the stamp overlay faded in proportionally to drag distance."
    ],
    [
        "Filtering System (Location & Price)",
        "Parameterized query testing",
        "Applied Province, City, and Adoption Fee filters individually and in combination. Verified that filtered results matched expected subsets by cross-referencing with direct database queries in the Supabase dashboard."
    ],
    [
        "Pet Data Storage & Retrieval",
        "Database integrity checks",
        "Inserted pet records via the Shelter Dashboard and verified column constraints, foreign key relationships, and RLS policies by attempting unauthorized reads/writes from different user sessions in the Supabase SQL editor."
    ],
    [
        "User Interaction Tracking (Like / Skip)",
        "Interaction log verification",
        "Performed a series of swipes and queried the interactions table directly to confirm each Like/Skip was recorded with the correct user ID, pet ID, and interaction type. Verified that previously swiped pets were excluded from subsequent browse sessions."
    ],
    [
        "Recommendation Algorithm (Preference-based Scoring)",
        "Score calculation validation",
        "Created test users with known preference profiles and manually computed expected match scores using the documented formula (species: 30, energy: 25, kid-friendly: 25, pet-friendly: 20). Compared manual calculations against system output to confirm accuracy."
    ],
    [
        "Behavioral Feedback Enhancement",
        "Incremental swipe pattern testing",
        "Performed controlled swipe sequences (e.g., consistently liking high-energy dogs) and verified that the behavioral bonus shifted rankings after the 5-swipe activation threshold. Confirmed bonus points matched the documented ratio formula."
    ],
    [
        "Explainable Recommendation Labels",
        "Label generation rule-tree testing",
        "Tested user-pet pairings with known high-scoring dimensions and verified that the correct natural language labels appeared on each card (e.g., \u201cHigh activity level match\u201d for matching energy levels). Confirmed labels disappeared when scores fell below threshold."
    ],
    [
        "Personalized Recommendation Display",
        "Sort order and UI verification",
        "Compared the displayed card order against the raw scored list from the recommendation engine. Verified that pets appeared in descending match percentage order and that ties were broken by listing recency (newer first)."
    ],
    [
        "Favorites / Saved Pets",
        "State synchronization testing",
        "Liked pets via the swipe interface and immediately checked the /favorites page. Verified real-time cache invalidation by confirming new favorites appeared without a manual page refresh, and that unliking removed them instantly."
    ],
    [
        "Chat / Messaging System",
        "Real-time WebSocket testing",
        "Opened two browser sessions (Adopter and Shelter) simultaneously. Sent messages from both sides and confirmed instant delivery via Supabase Realtime subscriptions. Verified message persistence, ordering by timestamp, and unread indicators."
    ],
    [
        "Adoption Fee Support",
        "Field validation and display testing",
        "Created pet listings with various fee values (zero, typical, high) and invalid inputs (negative, non-numeric). Confirmed Zod validation rejected invalid entries, valid fees displayed correctly on pet detail pages, and fee-based filtering returned accurate results."
    ],
    [
        "Frontend\u2013Backend Integration",
        "End-to-end data flow testing",
        "Traced full request lifecycles from UI action to Supabase query and back. Used browser DevTools Network tab to inspect API calls, verify response payloads, and confirm that authentication tokens were correctly attached to all protected requests."
    ],
    [
        "Shelter Dashboard Statistics",
        "Dashboard data accuracy checks",
        "Compared dashboard-displayed statistics against direct SQL aggregate queries on the database to verify that counts and summaries matched the actual data state."
    ],
    [
        "User Profile Editing",
        "Profile update and validation testing",
        "Updated profile fields and verified changes persisted across sessions. Tested edge cases such as clearing required fields and submitting partial updates to confirm validation behavior."
    ],
]

# Find the paragraph that contains "7. Validation Results Table" or "Validation Results Table"
insert_index = None
for i, para in enumerate(doc.paragraphs):
    if 'Validation Results Table' in para.text and ('7.' in para.text or para.text.strip().startswith('7')):
        insert_index = i
        break

if insert_index is None:
    print("ERROR: Could not find 'Validation Results Table' heading")
    exit(1)

print(f"Found Validation Results Table heading at paragraph index {insert_index}")

# We need to insert elements before the Validation Results heading.
# python-docx doesn't have a direct insert method, so we manipulate the XML.
from lxml import etree
from docx.oxml.ns import qn

body = doc.element.body
validation_heading_element = doc.paragraphs[insert_index]._element

# Create heading paragraph "Testing & Debugging Methodology"
heading_para = deepcopy(doc.paragraphs[insert_index]._element)
# Clear existing runs
for run in heading_para.findall(qn('w:r')):
    heading_para.remove(run)
# Also clear bookmark starts/ends
for bm in heading_para.findall(qn('w:bookmarkStart')):
    heading_para.remove(bm)
for bm in heading_para.findall(qn('w:bookmarkEnd')):
    heading_para.remove(bm)

# Add new run with the heading text
new_run = deepcopy(doc.paragraphs[insert_index]._element.findall(qn('w:r'))[0])
for t in new_run.findall(qn('w:t')):
    new_run.remove(t)
t_elem = etree.SubElement(new_run, qn('w:t'))
t_elem.text = 'Testing & Debugging Methodology'
heading_para.append(new_run)

# Insert heading before the validation results heading
body.insert(list(body).index(validation_heading_element), heading_para)

# Create intro paragraph
intro_text = "The following table documents the testing and debugging approach used to validate each functional requirement prior to final acceptance. Each requirement was verified through a combination of manual testing, data inspection, and cross-browser validation to ensure correctness and reliability."
intro_para_elem = etree.SubElement(body, qn('w:p'))  # temporary, we'll move it

# Copy paragraph properties from a body paragraph for consistent styling
# Find a normal body paragraph to copy style from
normal_para = None
for p in doc.paragraphs:
    if p.style and p.style.name and 'Normal' in p.style.name and p.text.strip():
        normal_para = p
        break

intro_para_elem = deepcopy(normal_para._element) if normal_para else etree.SubElement(body, qn('w:p'))
for run in intro_para_elem.findall(qn('w:r')):
    intro_para_elem.remove(run)
run_elem = etree.SubElement(intro_para_elem, qn('w:r'))
t_elem = etree.SubElement(run_elem, qn('w:t'))
t_elem.text = intro_text

# Remove the temporary element from end of body and insert before validation heading
if intro_para_elem in list(body):
    body.remove(intro_para_elem)
body.insert(list(body).index(validation_heading_element), intro_para_elem)

# Now create the table
# First, let's find the existing validation table to copy its style
existing_tables = doc.tables
ref_table = existing_tables[0] if existing_tables else None

# Build table with python-docx by adding at end then moving
table = doc.add_table(rows=len(testing_data) + 1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Style the table
if ref_table and ref_table.style:
    table.style = ref_table.style
else:
    table.style = 'Table Grid'

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(3.5)

# Header row
headers = ["Requirement", "Testing Method", "How It Was Tested / Debugged"]
for j, header in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(header)
    run.bold = True
    run.font.size = Pt(9)

# Data rows
for i, row_data in enumerate(testing_data):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i + 1].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(9)

# Move the table element before the validation heading
tbl_elem = table._tbl
body.remove(tbl_elem)
body.insert(list(body).index(validation_heading_element), tbl_elem)

# Add a blank paragraph between the table and the heading for spacing
spacer = etree.SubElement(body, qn('w:p'))
body.remove(spacer)
body.insert(list(body).index(validation_heading_element), spacer)

# Save
doc.save(doc_path)
print("Successfully inserted Testing & Debugging Methodology table!")
